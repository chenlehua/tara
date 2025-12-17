"""
Word Parser
===========

Parser for Word documents (.doc, .docx).
"""

import io
from typing import Any, Dict

from docx import Document as DocxDocument
from app.common.utils import get_logger

from .base_parser import BaseParser

logger = get_logger(__name__)


class WordParser(BaseParser):
    """Word document parser."""

    async def parse(
        self,
        content: bytes,
        filename: str,
        enable_ocr: bool = True,
        enable_structure: bool = True,
    ) -> Dict[str, Any]:
        """Parse Word document."""
        result = {
            "title": filename,
            "content": "",
            "structure": {},
            "metadata": {},
            "page_count": 0,
            "ocr_result": {},
        }

        try:
            doc_file = io.BytesIO(content)
            doc = DocxDocument(doc_file)

            # Extract metadata
            core_props = doc.core_properties
            result["metadata"] = {
                "author": core_props.author or "",
                "title": core_props.title or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
            }

            if core_props.title:
                result["title"] = core_props.title

            # Extract content
            paragraphs = []
            structure = {"sections": [], "tables": []}

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

                    # Detect headings for structure
                    if para.style and para.style.name.startswith("Heading"):
                        level = (
                            int(para.style.name.replace("Heading ", ""))
                            if "Heading " in para.style.name
                            else 1
                        )
                        structure["sections"].append(
                            {
                                "title": text,
                                "level": level,
                            }
                        )

            result["content"] = "\n\n".join(paragraphs)

            # Extract tables
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)

                structure["tables"].append(
                    {
                        "index": i,
                        "data": table_data,
                        "rows": len(table_data),
                        "cols": len(table_data[0]) if table_data else 0,
                    }
                )

            if enable_structure:
                result["structure"] = structure

        except Exception as e:
            logger.error(f"Failed to parse Word document: {e}")
            raise

        return result
