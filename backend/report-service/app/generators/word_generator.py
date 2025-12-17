"""Word document report generator."""

import io
from datetime import datetime
from typing import Any, Dict, List, Optional

import httpx
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from app.common.config import settings
from app.common.utils import get_logger

logger = get_logger(__name__)

# Diagram service URL
DIAGRAM_SERVICE_URL = getattr(settings, 'diagram_service_url', 'http://diagram-service:8005')


class WordGenerator:
    """Generate Word document reports."""

    def __init__(self):
        self._diagram_cache: Dict[str, bytes] = {}

    async def _fetch_diagram(
        self,
        diagram_type: str,
        project_id: int,
        threat_id: Optional[int] = None,
    ) -> Optional[io.BytesIO]:
        """Fetch a diagram from the diagram service."""
        try:
            async with httpx.AsyncClient(timeout=30.0) as client:
                if diagram_type == "attack_tree" and threat_id:
                    url = f"{DIAGRAM_SERVICE_URL}/api/v1/diagrams/attack-tree/{threat_id}?format=png"
                else:
                    url = f"{DIAGRAM_SERVICE_URL}/api/v1/diagrams/{diagram_type.replace('_', '-')}/{project_id}?format=png"
                
                response = await client.get(url)
                
                if response.status_code == 200:
                    return io.BytesIO(response.content)
                else:
                    logger.warning(f"Failed to fetch diagram {diagram_type}: {response.status_code}")
        except Exception as e:
            logger.warning(f"Error fetching diagram {diagram_type}: {e}")
        
        return None

    def _add_diagram_to_doc(
        self,
        doc: Document,
        diagram_buffer: Optional[io.BytesIO],
        caption: str,
        width: float = 6.0,  # inches
    ) -> None:
        """Add a diagram image to the document."""
        if diagram_buffer:
            try:
                diagram_buffer.seek(0)
                doc.add_picture(diagram_buffer, width=Inches(width))
                # Add caption
                caption_para = doc.add_paragraph()
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption_para.add_run(caption)
                caption_run.italic = True
                caption_run.font.size = Pt(10)
                doc.add_paragraph()  # Spacing
            except Exception as e:
                logger.warning(f"Failed to add diagram to Word doc: {e}")
                doc.add_paragraph(f"[图表加载失败: {caption}]", style="Intense Quote")

    def _set_cell_shading(self, cell, color: str):
        """Set cell background color."""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading)

    def _create_styled_table(self, doc: Document, headers: List[str], rows: List[List[str]]) -> any:
        """Create a styled table with headers."""
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            self._set_cell_shading(cell, "4F46E5")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(10)

        # Data rows
        for row_idx, row_data in enumerate(rows):
            row = table.rows[row_idx + 1]
            for col_idx, cell_text in enumerate(row_data):
                cell = row.cells[col_idx]
                cell.text = str(cell_text) if cell_text else "-"
                # Alternate row colors
                if row_idx % 2 == 0:
                    self._set_cell_shading(cell, "F8FAFC")

        return table

    async def generate(
        self,
        data: dict,
        template: str = "iso21434",
        sections: list[str] = None,
    ) -> io.BytesIO:
        """Generate Word document."""
        doc = Document()

        # Get content from data
        content = data.get("content", {})
        project = content.get("project", data.get("project", {}))

        # Set document properties
        core_properties = doc.core_properties
        core_properties.author = "TARA System"
        core_properties.title = f"TARA Report - {project.get('name', 'Unknown')}"

        # Title
        title = doc.add_heading("TARA Analysis Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Project info
        doc.add_paragraph()
        info_para = doc.add_paragraph()
        info_para.add_run("Project: ").bold = True
        info_para.add_run(project.get("name", "N/A"))

        info_para = doc.add_paragraph()
        info_para.add_run("Vehicle Type: ").bold = True
        info_para.add_run(project.get("vehicle_type", "N/A"))

        info_para = doc.add_paragraph()
        info_para.add_run("Standard: ").bold = True
        info_para.add_run(project.get("standard", "ISO/SAE 21434"))

        info_para = doc.add_paragraph()
        info_para.add_run("Generated: ").bold = True
        info_para.add_run(datetime.now().strftime("%Y-%m-%d %H:%M"))

        doc.add_paragraph()

        # Table of contents placeholder
        doc.add_heading("Table of Contents", 1)
        doc.add_paragraph("1. Executive Summary")
        doc.add_paragraph("2. Scope Definition")
        doc.add_paragraph("   2.1 Item Definition")
        doc.add_paragraph("   2.2 Item Boundary")
        doc.add_paragraph("   2.3 System Architecture")
        doc.add_paragraph("   2.4 Software Architecture")
        doc.add_paragraph("3. Asset Identification")
        doc.add_paragraph("4. Threat Analysis")
        doc.add_paragraph("5. Risk Assessment")
        doc.add_paragraph("6. Control Measures")
        doc.add_paragraph("Appendices")
        doc.add_page_break()

        # Generate sections
        if template == "iso21434":
            await self._generate_iso21434_sections(doc, data, sections)
        else:
            await self._generate_default_sections(doc, data)

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer

    async def _generate_iso21434_sections(
        self,
        doc: Document,
        data: dict,
        sections: list[str] = None,
    ):
        """Generate ISO 21434 format sections."""
        content = data.get("content", {})
        assets = content.get("assets", [])
        threats = content.get("threats", [])
        measures = content.get("control_measures", [])
        risk_dist = content.get("risk_distribution", {})
        
        high_risk = risk_dist.get("CAL-4", 0) + risk_dist.get("CAL-3", 0)

        # 1. Executive Summary
        doc.add_heading("1. Executive Summary", 1)
        doc.add_paragraph(
            "This document presents the results of the Threat Analysis and Risk Assessment "
            "(TARA) conducted in accordance with ISO/SAE 21434 cybersecurity engineering "
            "requirements for road vehicles."
        )

        doc.add_paragraph()
        doc.add_paragraph("Summary Statistics:", style="Intense Quote")

        stats_rows = [
            ["Total Assets", str(len(assets))],
            ["Total Threats", str(len(threats))],
            ["High Risk Items (CAL-3/4)", str(high_risk)],
            ["Control Measures", str(len(measures))],
        ]
        self._create_styled_table(doc, ["Metric", "Value"], stats_rows)

        doc.add_page_break()

        # 2. Scope Definition
        doc.add_heading("2. Scope Definition", 1)
        doc.add_paragraph(
            "The scope of this TARA covers the cybersecurity analysis of the vehicle "
            "electrical/electronic architecture and connected systems."
        )

        doc.add_heading("2.1 Item Definition", 2)
        project = content.get("project", data.get("project", {}))
        doc.add_paragraph(
            f"Target system: {project.get('name', 'Vehicle System')}"
        )
        doc.add_paragraph(
            f"Vehicle type: {project.get('vehicle_type', 'N/A')}"
        )
        doc.add_paragraph(
            f"Applicable standard: {project.get('standard', 'ISO/SAE 21434')}"
        )

        # Get project_id for diagrams
        project_id = content.get("project", {}).get("id") or data.get("project_id", 1)
        
        doc.add_heading("2.2 Item Boundary 项目边界", 2)
        doc.add_paragraph(
            "The item boundary defines the scope of this TARA analysis. "
            "Components within the boundary are subject to cybersecurity assessment."
        )
        
        # Fetch and add item boundary diagram
        item_boundary_diagram = await self._fetch_diagram("item-boundary", project_id)
        self._add_diagram_to_doc(doc, item_boundary_diagram, "图 2.1 项目边界图 (Item Boundary Diagram)")
        
        # Item Boundary Table
        boundary_rows = [
            ["Core Processing Unit 核心处理单元", "Internal", "Yes"],
            ["Security Module (HSM/SE) 安全模块", "Internal", "Yes"],
            ["Communication Interface 通信接口", "Internal", "Yes"],
            ["Data Storage 数据存储", "Internal", "Yes"],
            ["External Network 外部网络", "External", "Interface"],
            ["Cloud Services 云端服务", "External", "Interface"],
            ["User Interface 用户接口", "External", "Interface"],
        ]
        self._create_styled_table(doc, ["Component 组件", "Type 类型", "Within Boundary 边界内"], boundary_rows)
        
        doc.add_paragraph()
        doc.add_paragraph(
            f"Number of identified assets: {len(assets)}"
        )
        doc.add_paragraph(
            f"Number of interfaces analyzed: {sum(len(a.get('interfaces', [])) for a in assets)}"
        )

        doc.add_heading("2.3 System Architecture 系统架构", 2)
        doc.add_paragraph(
            "The system architecture consists of multiple layers following industry best practices. "
            "Each layer has specific security responsibilities and controls."
        )
        
        # Fetch and add system architecture diagram
        system_arch_diagram = await self._fetch_diagram("system-architecture", project_id)
        self._add_diagram_to_doc(doc, system_arch_diagram, "图 2.2 系统架构图 (System Architecture Diagram)")
        
        # System Architecture Table
        arch_rows = [
            ["Application Layer 应用层", "HMI, Navigation, Media, ADAS", "Input Validation, Access Control"],
            ["Service Layer 服务层", "Security, Comm, Diagnostic, OTA", "Authentication, Encryption"],
            ["Middleware Layer 中间件层", "AUTOSAR, Hypervisor, OS, Crypto", "Isolation, Secure Boot"],
            ["Hardware Layer 硬件层", "SoC/MCU, HSM/SE, Memory, Network", "Hardware Security, Key Storage"],
        ]
        self._create_styled_table(doc, ["Layer 层级", "Components 组件", "Security Focus 安全重点"], arch_rows)

        doc.add_heading("2.4 Software Architecture 软件架构", 2)
        doc.add_paragraph(
            "Key software modules and their security responsibilities are detailed below. "
            "The software architecture follows a modular design with clear security boundaries."
        )
        
        # Fetch and add software architecture diagram
        sw_arch_diagram = await self._fetch_diagram("software-architecture", project_id)
        self._add_diagram_to_doc(doc, sw_arch_diagram, "图 2.3 软件架构图 (Software Architecture Diagram)")
        
        # Software Architecture Table
        sw_rows = [
            ["Security Manager 安全管理器", "Central security control", "Policy enforcement, Key management"],
            ["Application Manager 应用管理器", "App lifecycle management", "Sandbox, Privilege control"],
            ["Communication Manager 通信管理器", "Network protocols", "TLS/DTLS, Certificate validation"],
            ["Update Manager OTA升级管理器", "OTA updates", "Signature verification, Rollback"],
            ["Crypto Library 加密库", "Cryptographic operations", "AES, RSA, ECC, Hash functions"],
            ["Key Management 密钥管理", "Key lifecycle", "HSM integration, Key derivation"],
        ]
        self._create_styled_table(doc, ["Module 模块", "Function 功能", "Security Measures 安全措施"], sw_rows)

        doc.add_page_break()

        # 3. Asset Identification
        doc.add_heading("3. Asset Identification", 1)
        doc.add_paragraph(
            "The following assets have been identified and categorized according to their "
            "cybersecurity relevance."
        )
        
        # Fetch and add asset graph diagram
        doc.add_heading("3.1 Asset Relationship Graph 资产关系图", 2)
        asset_graph_diagram = await self._fetch_diagram("asset-graph", project_id)
        self._add_diagram_to_doc(doc, asset_graph_diagram, "图 3.1 资产关系图 (Asset Relationship Graph)")
        
        # Fetch and add data flow diagram
        doc.add_heading("3.2 Data Flow Diagram 数据流图", 2)
        data_flow_diagram = await self._fetch_diagram("data-flow", project_id)
        self._add_diagram_to_doc(doc, data_flow_diagram, "图 3.2 数据流图 (Data Flow Diagram)")
        
        doc.add_heading("3.3 Asset List 资产清单", 2)

        if assets:
            asset_rows = []
            for asset in assets[:30]:  # Limit to 30
                interfaces = asset.get("interfaces", [])
                iface_str = ", ".join([
                    i.get("type", str(i)) if isinstance(i, dict) else str(i) 
                    for i in interfaces[:3]
                ])
                if len(interfaces) > 3:
                    iface_str += f" (+{len(interfaces) - 3})"
                
                asset_rows.append([
                    asset.get("name", "N/A")[:40],
                    asset.get("type", asset.get("asset_type", "N/A")),
                    asset.get("security_level", asset.get("criticality", "CAL-2")),
                    iface_str or "-",
                ])
            
            self._create_styled_table(
                doc, 
                ["Asset Name", "Type", "Security Level", "Interfaces"], 
                asset_rows
            )
            
            if len(assets) > 30:
                doc.add_paragraph(f"... and {len(assets) - 30} more assets")
        else:
            doc.add_paragraph("No assets identified.", style="Intense Quote")

        doc.add_page_break()

        # 4. Threat Analysis
        doc.add_heading("4. Threat Analysis", 1)
        doc.add_paragraph(
            "Threats were systematically identified using the STRIDE methodology "
            "(Spoofing, Tampering, Repudiation, Information Disclosure, "
            "Denial of Service, Elevation of Privilege)."
        )
        
        # Fetch and add attack tree diagram
        if threats:
            first_threat = threats[0]
            first_threat_id = first_threat.get("id")
            if first_threat_id and isinstance(first_threat_id, int):
                doc.add_heading("4.1 Attack Tree Analysis 攻击树分析", 2)
                attack_tree_diagram = await self._fetch_diagram("attack_tree", project_id, threat_id=first_threat_id)
                self._add_diagram_to_doc(doc, attack_tree_diagram, "图 4.1 攻击树示例 (Attack Tree Example)")
        
        doc.add_heading("4.2 Threat List 威胁清单", 2)

        if threats:
            threat_rows = []
            for threat in threats[:30]:  # Limit to 30
                threat_rows.append([
                    threat.get("id", "N/A"),
                    threat.get("name", "N/A")[:40],
                    threat.get("category_name", threat.get("category", "N/A")),
                    threat.get("risk_level", "CAL-2"),
                ])
            
            self._create_styled_table(
                doc, 
                ["ID", "Threat Name", "STRIDE Category", "Risk Level"], 
                threat_rows
            )
            
            if len(threats) > 30:
                doc.add_paragraph(f"... and {len(threats) - 30} more threats")
        else:
            doc.add_paragraph("No threats identified.", style="Intense Quote")

        doc.add_page_break()

        # 5. Risk Assessment
        doc.add_heading("5. Risk Assessment", 1)
        doc.add_paragraph(
            "Risk levels were determined based on the combination of impact severity "
            "and attack feasibility, in accordance with ISO/SAE 21434 methodology."
        )
        
        # Fetch and add risk matrix diagram
        doc.add_heading("5.1 Risk Matrix 风险矩阵", 2)
        risk_matrix_diagram = await self._fetch_diagram("risk-matrix", project_id)
        self._add_diagram_to_doc(doc, risk_matrix_diagram, "图 5.1 风险矩阵 (Risk Matrix - ISO/SAE 21434)")

        doc.add_heading("5.2 Risk Distribution 风险分布", 2)
        if risk_dist:
            risk_rows = [
                ["CAL-4 (Critical)", str(risk_dist.get("CAL-4", 0)), "Immediate action required"],
                ["CAL-3 (High)", str(risk_dist.get("CAL-3", 0)), "Priority treatment needed"],
                ["CAL-2 (Medium)", str(risk_dist.get("CAL-2", 0)), "Planned treatment"],
                ["CAL-1 (Low)", str(risk_dist.get("CAL-1", 0)), "Acceptable risk"],
            ]
            self._create_styled_table(doc, ["Risk Level", "Count", "Action Required"], risk_rows)

        doc.add_heading("5.2 Impact Assessment", 2)
        doc.add_paragraph(
            "Impact assessment considers safety, financial, operational, and privacy impacts "
            "as defined in ISO/SAE 21434."
        )

        doc.add_heading("5.3 Attack Feasibility Assessment", 2)
        doc.add_paragraph(
            "Attack feasibility is evaluated based on attack potential factors including "
            "elapsed time, expertise, knowledge of item, window of opportunity, and equipment."
        )

        doc.add_page_break()

        # 6. Control Measures
        doc.add_heading("6. Control Measures", 1)
        doc.add_paragraph(
            "Based on the risk assessment results, the following security control measures "
            "have been recommended."
        )

        if measures:
            measure_rows = []
            for measure in measures[:30]:  # Limit to 30
                eff = measure.get("effectiveness", "medium")
                eff_text = {"high": "High", "medium": "Medium", "low": "Low"}.get(eff, "Medium")
                
                measure_rows.append([
                    measure.get("name", "N/A")[:40],
                    measure.get("control_type", measure.get("category", "preventive")),
                    eff_text,
                    measure.get("iso21434_ref", "-"),
                ])
            
            self._create_styled_table(
                doc, 
                ["Measure Name", "Type", "Effectiveness", "ISO 21434 Ref"], 
                measure_rows
            )
            
            if len(measures) > 30:
                doc.add_paragraph(f"... and {len(measures) - 30} more measures")
        else:
            doc.add_paragraph("No control measures defined.", style="Intense Quote")

        doc.add_page_break()

        # Appendices
        doc.add_heading("Appendices", 1)
        doc.add_heading("A. Glossary", 2)
        doc.add_paragraph("TARA - Threat Analysis and Risk Assessment")
        doc.add_paragraph(
            "STRIDE - Spoofing, Tampering, Repudiation, Information Disclosure, Denial of Service, Elevation of Privilege"
        )
        doc.add_paragraph("CAL - Cybersecurity Assurance Level")
        doc.add_paragraph("ISO 21434 - Road vehicles — Cybersecurity engineering")

    async def _generate_default_sections(self, doc: Document, data: dict):
        """Generate default sections."""
        await self._generate_iso21434_sections(doc, data, None)
